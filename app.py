# --- EXTERNAL LIBRARIES ---
import streamlit as st # type: ignore
import pandas as pd # type: ignore
import numpy as np # type: ignore
import cv2 # type: ignore
from PIL import Image # type: ignore
import openpyxl # type: ignore
from openpyxl.styles import Font # type: ignore
from openpyxl.utils import get_column_letter # type: ignore
from docx import Document # type: ignore
from docx.shared import Inches # type: ignore

# --- PYTHON STANDARD LIBRARY ---
import os
import io
import re
import time
import zipfile
import shutil
from io import BytesIO
from datetime import datetime, date
import base64

# --- PAGE CONFIGURATION MUST BE AT THE VERY TOP ---
st.set_page_config(page_title="Workspace Tools", page_icon="🧰", layout="wide")

# ========== SESSION STATE INIT ==========
if 'entered' not in st.session_state:
    st.session_state.entered = False  # Show welcome page initially

# ========== WELCOME PAGE ==========
if not st.session_state.entered:
    # Hide sidebar and set full-width welcome layout
    st.markdown(
        """
        <style>
        section[data-testid="stSidebar"] {
            display: none !important;
        }
        .main > div {
            padding: 2rem 4rem;
        }
        .welcome-card {
            background: white;
            border-radius: 30px;
            padding: 3rem 4rem;
            box-shadow: 0 20px 60px rgba(0,0,0,0.08);
            text-align: center;
            max-width: 700px;
            margin: 5rem auto;
            border: 1px solid rgba(255,255,255,0.3);
        }
        .welcome-title {
            font-size: 3.2rem;
            font-weight: 700;
            color: #1e2a3a;
            margin-bottom: 0.5rem;
        }
        .welcome-sub {
            font-size: 1.2rem;
            color: #4a5a6a;
            margin-bottom: 2rem;
        }
        .welcome-emoji {
            font-size: 4rem;
            margin-bottom: 1rem;
        }
        .stButton > button {
            background: #2e7d32 !important;
            color: white !important;
            border: none !important;
            padding: 0.8rem 3rem !important;
            border-radius: 40px !important;
            font-size: 1.2rem !important;
            font-weight: 600 !important;
            box-shadow: 0 8px 20px rgba(46, 125, 50, 0.25);
            transition: all 0.3s ease;
        }
        .stButton > button:hover {
            transform: translateY(-3px);
            box-shadow: 0 12px 28px rgba(46, 125, 50, 0.35);
            background: #1b5e20 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("""
        <div class="welcome-card">
            <div class="welcome-emoji">🧰</div>
            <div class="welcome-title">Welcome to Workspace Tools</div>
            <div class="welcome-sub">
                Your all‑in‑one toolkit for VRP mapping, field results, and signature extraction.<br>
                Click the button below to get started.
            </div>
        </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("🚀 Enter Workspace", use_container_width=False):
            st.session_state.entered = True
            st.rerun()
    st.stop()


# ========== MINIMAL STYLING (button animations & dropdown visibility) ==========
st.markdown(
    """
    <style>
    /* Button hover/active animations (original) */
    .stButton > button, .stDownloadButton > button {
        transition: all 0.3s cubic-bezier(0.25, 0.8, 0.25, 1);
        border-radius: 8px !important;
    }
    .stButton > button:hover, .stDownloadButton > button:hover {
        transform: translateY(-4px);
        box-shadow: 0 10px 20px rgba(0, 0, 0, 0.2) !important;
        border-color: #2e7d32 !important;
        animation: gentle-pulse 1.5s infinite ease-in-out;
        z-index: 1;
    }
    .stButton > button:active, .stDownloadButton > button:active {
        transform: translateY(2px) scale(0.98) !important;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1) !important;
        animation: none;
    }
    @keyframes gentle-pulse {
        0% { transform: translateY(-4px) scale(1); }
        50% { transform: translateY(-4px) scale(1.03); }
        100% { transform: translateY(-4px) scale(1); }
    }

    /* Dropdown & text area – white background for readability */
    .stSelectbox div[data-baseweb="select"] {
        background-color: #ffffff !important;
        border-radius: 8px;
        border: 1px solid #d0d9e8;
    }
    .stSelectbox div[data-baseweb="select"] > div {
        background-color: #ffffff !important;
        color: #1e2a3a !important;
    }
    .stSelectbox ul {
        background-color: #ffffff !important;
    }
    .stSelectbox li {
        color: #1e2a3a !important;
    }
    .stTextArea textarea {
        background-color: #ffffff !important;
        color: #1e2a3a !important;
        border-radius: 8px;
        border: 1px solid #d0d9e8;
    }
    .stTextArea textarea:focus {
        border-color: #2e7d32;
        box-shadow: 0 0 0 2px rgba(46, 125, 50, 0.2);
    }

    /* Sidebar subtle style (optional) */
    .stSidebar {
        background: #ffffffdd;
        backdrop-filter: blur(4px);
        border-right: 1px solid rgba(0,0,0,0.05);
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ----- Sound notification helper -----
def play_completion_sound():
    """Inject HTML/JS to play a short chime sound."""
    html = f"""
    <audio id="completion-sound" style="display:none;"></audio>
    <script>
    (function(){{
        var audioCtx = new (window.AudioContext || window.webkitAudioContext)();
        var oscillator = audioCtx.createOscillator();
        var gainNode = audioCtx.createGain();
        oscillator.connect(gainNode);
        gainNode.connect(audioCtx.destination);
        oscillator.frequency.value = 880;
        oscillator.type = 'sine';
        gainNode.gain.value = 0.3;
        oscillator.start();
        gainNode.gain.exponentialRampToValueAtTime(0.0001, audioCtx.currentTime + 0.3);
        oscillator.stop(audioCtx.currentTime + 0.3);
    }})();
    </script>
    """
    st.components.v1.html(html, height=0)

# ========== SESSION STATE (other states) ==========
if 'uploader_key' not in st.session_state:
    st.session_state.uploader_key = 0
if 'processed_data' not in st.session_state:
    st.session_state.processed_data = None
if 'release_data' not in st.session_state:
    st.session_state.release_data = None
if 'generated_filename' not in st.session_state:
    st.session_state.generated_filename = ""
if 'release_filename' not in st.session_state:
    st.session_state.release_filename = ""
if 'selected_type' not in st.session_state:
    st.session_state.selected_type = None
if 'show_account_type_modal' not in st.session_state:
    st.session_state.show_account_type_modal = False
if 'process_confirm' not in st.session_state:
    st.session_state.process_confirm = False
if 'is_multiple_files' not in st.session_state:
    st.session_state.is_multiple_files = False
if 'cms_id_warning' not in st.session_state:
    st.session_state.cms_id_warning = None
if 'version' not in st.session_state:
    st.session_state.version = "MC2"

# Field Result session state
if 'field_result_buffer' not in st.session_state:
    st.session_state.field_result_buffer = None
if 'field_result_filename' not in st.session_state:
    st.session_state.field_result_filename = ""
if 'field_result_processed' not in st.session_state:
    st.session_state.field_result_processed = False
if 'field_result_elapsed' not in st.session_state:
    st.session_state.field_result_elapsed = 0.0

def reset_app():
    st.session_state.uploader_key += 1
    st.session_state.processed_data = None
    st.session_state.release_data = None
    st.session_state.selected_type = None
    st.session_state.show_account_type_modal = False
    st.session_state.process_confirm = False
    st.session_state.is_multiple_files = False
    st.session_state.cms_id_warning = None
    st.rerun()

# ========== SIDEBAR ==========
st.sidebar.title("🛠️ Workspace")
selected_tool = st.sidebar.radio(
    "Select Tool:",
    ["VRP Mapper", "Field Result", "E-SIGN FIXER"]
)

st.sidebar.divider()
st.sidebar.info("Select a tool from the menu above to get started.")


# ==========================================
# TOOL 1: VRP ACCOUNT TEMPLATE MAPPER
# ==========================================
if selected_tool == "VRP Mapper":
    st.title("🚀 VRP Mapper")

    # --- Version Dropdown ---
    version = st.selectbox(
        "Select Version:",
        options=["MC2", "FCL", "OTS"],
        index=["MC2", "FCL", "OTS"].index(st.session_state.version),
        key="version_select"
    )
    st.session_state.version = version

    # --- 1. SETTINGS & TEMPLATE ---
    template_filename = 'demand_letter_template.csv'
    if not os.path.exists(template_filename):
        st.error(f"❌ Error: '{template_filename}' not found.")
        st.stop()
    df_tmp = pd.read_csv(template_filename, dtype=str)

    # --- CMS ID REFERENCE LOADING ---
    @st.cache_data
    def load_cms_data():
        ref_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "cmd_id.xlsx")
        if os.path.exists(ref_path):
            try:
                df_cms = pd.read_excel(ref_path, dtype=str)
                df_cms.columns = df_cms.columns.str.strip().str.upper()
                cms_col = next((c for c in df_cms.columns if c in ['CMS ID', 'CMS_ID']), None)
                ch_col = next((c for c in df_cms.columns if c in ['CH CODE', 'CH_CODE']), None)
                if ch_col and cms_col:
                    return dict(zip(df_cms[ch_col].str.strip().str.upper(), df_cms[cms_col].str.strip()))
            except Exception:
                pass
        return {}

    cms_mapping = load_cms_data()

    # --- AMOUNTS REFERENCE LOADING ---
    @st.cache_data
    def load_amounts_data():
        ref_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "amounts.csv")
        if os.path.exists(ref_path):
            try:
                df_amt = pd.read_csv(ref_path, dtype=str)
                df_amt.columns = df_amt.columns.str.strip()
                ch_col = next((c for c in df_amt.columns if c.upper().replace(" ", "") in ['CHCODE', 'CH_CODE']), None)
                ob_col = next((c for c in df_amt.columns if c.strip().upper() == 'AMOUNT OB'), None)
                due_col = next((c for c in df_amt.columns if c.strip().upper() == 'PRINCIPAL AMOUNT DUE'), None)
                if ch_col and ob_col and due_col:
                    df_amt[ch_col] = df_amt[ch_col].str.strip().str.upper()
                    ob_dict = dict(zip(df_amt[ch_col], df_amt[ob_col].str.strip().str.replace(',', '', regex=False)))
                    due_dict = dict(zip(df_amt[ch_col], df_amt[due_col].str.strip().str.replace(',', '', regex=False)))
                    return ob_dict, due_dict
            except Exception:
                pass
        return {}, {}

    amounts_ob_mapping, amounts_due_mapping = load_amounts_data()

    # --- AREA CLUSTER MAPPING ---
    AREA_MAPPING = {
        'ABRA': 'NORTH LUZON', 'AURORA': 'NORTH LUZON', 'BATAAN': 'NORTH LUZON',
        'BENGUET': 'NORTH LUZON', 'BULACAN': 'NORTH LUZON', 'CAGAYAN': 'NORTH LUZON',
        'ILOCOS NORTE': 'NORTH LUZON', 'ILOCOS SUR': 'NORTH LUZON', 'ISABELA': 'NORTH LUZON',
        'LA UNION': 'NORTH LUZON', 'NUEVA ECIJA': 'NORTH LUZON', 'PAMPANGA': 'NORTH LUZON',
        'PANGASINAN': 'NORTH LUZON', 'TARLAC': 'NORTH LUZON', 'ZAMBALES': 'NORTH LUZON',
        'NCR': 'NCR', 'METRO MANILA': 'NCR',
        'BATANGAS': 'SOUTH LUZON', 'CAMARINES SUR': 'SOUTH LUZON', 'CAVITE': 'SOUTH LUZON',
        'LAGUNA': 'SOUTH LUZON', 'OCCIDENTAL MINDORO': 'SOUTH LUZON', 'ORIENTAL MINDORO': 'SOUTH LUZON',
        'QUEZON PROVINCE': 'SOUTH LUZON', 'RIZAL': 'SOUTH LUZON',
        'AKLAN': 'VISAYAS', 'ANTIQUE': 'VISAYAS', 'BOHOL': 'VISAYAS', 'CAPIZ': 'VISAYAS',
        'CEBU': 'VISAYAS', 'GUIMARAS': 'VISAYAS', 'ILOILO': 'VISAYAS', 'LEYTE': 'VISAYAS',
        'NEGROS OCCIDENTAL': 'VISAYAS', 'NEGROS ORIENTAL': 'VISAYAS',
        'AGUSAN DEL SUR': 'MINDANAO', 'BUKIDNON': 'MINDANAO', 'COTABATO': 'MINDANAO',
        'DAVAO DE ORO': 'MINDANAO', 'DAVAO DEL NORTE': 'MINDANAO', 'DAVAO DEL SUR': 'MINDANAO',
        'DAVAO ORIENTAL': 'MINDANAO', 'LANAO DEL NORTE': 'MINDANAO', 'LANAO DEL SUR': 'MINDANAO',
        'MAGUINDANAO': 'MINDANAO', 'MISAMIS OCCIDENTAL': 'MINDANAO', 'MISAMIS ORIENTAL': 'MINDANAO',
        'SARANGANI': 'MINDANAO', 'SOUTH COTABATO': 'MINDANAO', 'SULTAN KUDARAT': 'MINDANAO',
        'SULU': 'MINDANAO', 'SURIGAO DEL SUR': 'MINDANAO', 'TARLAC': 'MINDANAO', 'TAWI-TAWI': 'MINDANAO',
        'ZAMBOANGA DEL NORTE': 'MINDANAO', 'ZAMBOANGA DEL SUR': 'MINDANAO', 'ZAMBOANGA SIBUGAY': 'MINDANAO'
    }

    st.divider()

    # --- FILE UPLOADER based on version ---
    if version == "MC2":
        src_files = st.file_uploader(
            "Upload VRP ACCOUNTS CSV(s)",
            type=['csv'],
            accept_multiple_files=True,
            key=f"uploader_{st.session_state.uploader_key}"
        )
    else:  # FCL or OTS
        src_files = st.file_uploader(
            "Upload VRP ACCOUNTS Excel file",
            type=['xlsx'],
            accept_multiple_files=False,
            key=f"uploader_{st.session_state.uploader_key}"
        )

    # Optional filter (pasted codes) - works for all versions
    pasted_codes = st.text_area("Paste Reference Codes (Optional filter - one per line):", height=150)

    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("✨ Process and Map Data", use_container_width=True):
            st.session_state.process_confirm = True
    with col2:
        if st.button("🔄 Reset / Clear All", use_container_width=True):
            reset_app()

    if st.session_state.process_confirm:
        # --- READ DATA ---
        if version == "MC2":
            if not src_files:
                st.warning("Please upload at least one CSV file.")
                st.stop()
            df_list = []
            for file in src_files:
                df_temp = pd.read_csv(file, dtype=str, keep_default_na=False)
                df_temp.columns = df_temp.columns.str.strip()
                raw_name = os.path.splitext(file.name)[0].strip().upper()
                file_type = "DL"
                if "REVISIT" in raw_name and "NO DL" in raw_name:
                    file_type = "Trans/Details"
                elif "WITH DL" in raw_name:
                    file_type = "DL"
                elif "MC2" in raw_name and "OTHERS" in raw_name:
                    file_type = "DL"
                df_temp['_FILE_ASSIGNED_TYPE'] = file_type
                df_temp['_FILE_RAW_NAME'] = raw_name          # <--- NEW: store filename
                df_list.append(df_temp)
            df_master = pd.concat(df_list, ignore_index=True)
            is_multiple = True
        else:
            # FCL or OTS – single Excel file, read sheet "SUMMARY"
            if src_files is None:
                st.warning("Please upload an Excel file.")
                st.stop()
            try:
                xl = pd.ExcelFile(src_files)
                if "SUMMARY" not in xl.sheet_names:
                    st.error("❌ The Excel file does not contain a sheet named 'SUMMARY'. Please ensure the sheet name is exactly 'SUMMARY'.")
                    st.stop()
                df_master = pd.read_excel(src_files, sheet_name="SUMMARY", dtype=str, keep_default_na=False)
            except Exception as e:
                st.error(f"❌ Error reading sheet 'SUMMARY': {str(e)}")
                st.stop()
            df_master.columns = df_master.columns.str.strip()
            is_multiple = False

        start_time = time.time()
        progress_bar = st.progress(0, text="Initializing processing...")
        time.sleep(0.2)

        # --- FILTER by reference codes ---
        progress_bar.progress(15, text="Filtering target codes...")
        if pasted_codes.strip():
            codes_list = [c.strip() for c in pasted_codes.replace(',', '\n').split('\n') if c.strip()]
            ref_col = None
            for col in df_master.columns:
                col_clean = col.strip().upper().replace(" ", "").replace("_", "")
                if col_clean in ['REFCODE', 'REFERENCECODE', 'REFCODE']:
                    ref_col = col
                    break
            if ref_col:
                df_src = df_master[df_master[ref_col].isin(codes_list)].copy()
            else:
                st.error("Reference code column not found in the data.")
                st.stop()
        else:
            df_src = df_master.copy()

        df_src = df_src.reset_index(drop=True)
        df_out = pd.DataFrame(columns=df_tmp.columns, index=range(len(df_src)))
        time.sleep(0.2)

        # --- ROBUST COLUMN MAPPING ---
        progress_bar.progress(35, text="Mapping source columns to template...")
        col_map = {
            'account_no': ['ACCOUNT NUMBER', 'ACCOUNT_NO', 'ACCOUNTNUMBER'],
            'bank': ['BANK'],
            'placement': ['PLACEMENT'],
            'ch_code': ['CH CODE', 'CH_CODE', 'CHCODE'],
            'ch_name': ['CH NAME', 'CH_NAME', 'CHNAME'],
            'address_type': ['ADD TYPE', 'ADD_TYPE', 'ADDTYPE'],
            'address': ['ADDRESS'],
            'final_area': ['FINAL AREA', 'FINAL_AREA', 'FINALAREA'],
            'area': ['AREA'],
            'municipality': ['MUNICIPALITY'],
            'dl_type': ['DL TYPE', 'DL_TYPE', 'DLTYPE'],
            'ref_code': ['REFERENCE CODE', 'REF CODE', 'REF_CODE', 'REFCODE'],
            'autofield_date': ['AUTOFIELD DATE', 'AUTOFIELD_DATE', 'AUTOFIELDDATE'],
            'outstanding_balance': ['OB/PRINCIPAL', 'OB_PRINCIPAL', 'OBPRINCIPAL', 'OB'],
            'endorsement_date': ['ENDO DATE', 'ENDO_DATE', 'ENDODATE', 'ENDORSEMENT DATE'],
            'pullout_date': ['POUT DATE', 'POUT_DATE', 'POUTDATE', 'PULLOUT DATE'],
        }

        def find_column(source_cols, aliases):
            for col in source_cols:
                col_clean = col.strip().upper().replace(" ", "").replace("_", "").replace("/", "")
                for alias in aliases:
                    alias_clean = alias.strip().upper().replace(" ", "").replace("_", "").replace("/", "")
                    if col_clean == alias_clean:
                        return col
            return None

        matched_cols = {}
        for template_col, aliases in col_map.items():
            matched = find_column(df_src.columns, aliases)
            if matched:
                matched_cols[template_col] = matched

        for template_col, src_col in matched_cols.items():
            if template_col == 'account_no':
                df_out[template_col] = df_src[src_col].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()
                df_out[template_col] = df_out[template_col].replace(['nan', 'None', ''], '0')
            elif template_col == 'address':
                df_out[template_col] = (
                    df_src[src_col]
                    .fillna('')
                    .str.replace('Blk', 'Block', case=False, regex=False)
                    .str.replace('BRGY', 'BARANGAY', case=False, regex=False)
                    .str.replace(r'\bSTA\b\.?', 'SANTA', case=False, regex=True)
                    .str.replace(r'\bSTO\b\.?', 'SANTO', case=False, regex=True)
                    .str.upper()
                )
            elif template_col == 'outstanding_balance':
                vals = df_src[src_col].astype(str).str.strip().str.replace(',', '', regex=False)
                def format_general_number(val):
                    if val in ['nan', 'None', '', '0', '0.0']:
                        return '0'
                    try:
                        num = float(val)
                        if num == int(num):
                            return str(int(num))
                        else:
                            formatted = f"{num:.2f}"
                            formatted = formatted.rstrip('0').rstrip('.') if '.' in formatted else formatted
                            return formatted
                    except ValueError:
                        return val
                df_out[template_col] = vals.apply(format_general_number)
            else:
                df_out[template_col] = df_src[src_col]

        if 'final_area' not in matched_cols and 'area' in matched_cols:
            df_out['final_area'] = df_out['area']
        elif 'final_area' not in matched_cols and 'area' not in matched_cols:
            for col in df_src.columns:
                if col.strip().upper() == 'AREA':
                    df_out['final_area'] = df_src[col]
                    break

        time.sleep(0.2)

        # --- AMOUNT DUE & CMS ID ---
        progress_bar.progress(60, text="Calculating constants, amounts, and CMS IDs...")
        is_pif_homeloan = False
        if version == "OTS":
            is_pif_homeloan = True
        else:
            bank_col = None
            for col in df_src.columns:
                col_clean = col.strip().upper().replace(" ", "").replace("_", "").replace("/", "")
                if col_clean == 'BANK':
                    bank_col = col
                    break
            if bank_col:
                unique_banks = df_src[bank_col].astype(str).str.strip().str.upper().replace({'NAN': ''}).unique()
                is_pif_homeloan = any('PIF HOMELOAN' in bank for bank in unique_banks)

        if 'ch_code' in matched_cols:
            ch_codes = df_out['ch_code'].astype(str).str.strip().str.upper()
            df_out['amount_due'] = ch_codes.map(amounts_due_mapping).fillna('0')
            df_out['cms_id'] = ch_codes.map(cms_mapping).fillna('')
            if is_pif_homeloan:
                blank_cms_mask = df_out['cms_id'] == ''
                if blank_cms_mask.any():
                    missing_ch_codes = sorted(set(df_out.loc[blank_cms_mask, 'ch_code'].astype(str).str.strip()))
                    progress_bar.empty()
                    st.session_state.process_confirm = False
                    st.error(
                        f"🚨 **PROCESS HALTED:** Found {blank_cms_mask.sum()} PIF HOMELOAN row(s) with a blank CMS ID.\n\n"
                        f"**Missing CH CODE(s):** {', '.join(missing_ch_codes)}\n\n"
                        f"Please update your 'cmd_id.xlsx' reference file and try again."
                    )
                    st.stop()
        else:
            df_out['amount_due'] = '0'
            df_out['cms_id'] = ''

        df_out['shared_or_exclusive'] = "SHARED"

        # --- TYPE OF ACCOUNT & VISIT TYPE ---
        if version == "MC2":
            if is_multiple:
                df_out['type_of_account'] = df_src['_FILE_ASSIGNED_TYPE']
            else:
                if '_FILE_ASSIGNED_TYPE' in df_src.columns:
                    df_out['type_of_account'] = df_src['_FILE_ASSIGNED_TYPE']
                else:
                    df_out['type_of_account'] = "DL"
        elif version == "FCL":
            df_out['type_of_account'] = "DL"
        elif version == "OTS":
            df_out['type_of_account'] = "Trans/Details"

        if version == "MC2":
            def determine_visit_type_mc2(idx, row):
                # Force REGULAR for specific file patterns
                file_name = str(row.get('_FILE_RAW_NAME', '')).upper()
                if any(pattern in file_name for pattern in ["TAG_PIF WITH DL", "TAG_PIF REVISIT- NO DL", "TAG_MC2- OTHERS"]):
                    return "REGULAR"
                
                current_type = df_out.at[idx, 'type_of_account']
                if current_type == "Trans/Details":
                    return "REGULAR"
                remark = str(row.get('REMARKS', '')).strip().upper() if 'REMARKS' in row else ''
                bank = str(row.get('BANK', '')).strip().upper() if 'BANK' in row else ''
                if "PIF FORECLOSURE" in bank:
                    return "OTS"
                regular_banks = ["CBS HOUSING LOAN", "PIF PROVIDENT", "SBF MORTGAGE", "UBP HOME MORTGAGE", "SBC HOME LOAN", "SBF HOMELOAN"]
                if any(target in bank for target in regular_banks):
                    return "REGULAR"
                if remark in ["NEW ENDO", "REVISIT", "NEW ENDO-PRI ADD"]:
                    return "REGULAR"
                if remark == "CARAVAN":
                    return "CARAVAN"
                return ""
            df_out['visit_type'] = [determine_visit_type_mc2(idx, row) for idx, row in df_src.iterrows()]
        elif version == "FCL":
            df_out['visit_type'] = "OTS"
        elif version == "OTS":
            df_out['visit_type'] = "REGULAR"

        df_out['month'] = datetime.now().strftime('%B').upper()
        df_out['account_type'] = "HOUSING"
        df_out['form_code'] = "vid04qNT"

        if 'area' in matched_cols:
            df_out['area_cluster'] = df_src[matched_cols['area']].str.strip().str.upper().map(AREA_MAPPING)
        else:
            df_out['area_cluster'] = ""

        progress_bar.progress(80, text="Formatting dates...")
        date_fields = ['autofield_date', 'endorsement_date', 'pullout_date']
        for field in date_fields:
            if field in df_out.columns:
                df_out[field] = pd.to_datetime(df_out[field], errors='coerce').dt.strftime('%d-%m-%Y')
                df_out[field] = df_out[field].fillna('')
        time.sleep(0.2)

        progress_bar.progress(95, text="Generating final files...")
        total_accounts = len(df_out)
        if version == "MC2":
            bank_col = None
            for col in df_src.columns:
                col_clean = col.strip().upper().replace(" ", "").replace("_", "").replace("/", "")
                if col_clean == 'BANK':
                    bank_col = col
                    break
            if bank_col:
                unique_banks = df_src[bank_col].astype(str).str.strip().str.upper().replace({'NAN': ''}).unique()
                unique_banks = [b for b in unique_banks if b]
                if len(unique_banks) == 1 and unique_banks[0] == "PIF HOMELOAN":
                    bank_val = "PIF HOME LOAN"
                elif len(unique_banks) == 1:
                    bank_val = unique_banks[0]
                else:
                    bank_val = "MC2 OTHERS"
            else:
                bank_val = "FILTERED"
        elif version == "FCL":
            bank_val = "FCL"
        else:
            bank_val = "OTS"

        st.session_state.generated_filename = f"{bank_val}_{total_accounts}.csv"
        st.session_state.release_filename = f"Released_to_{bank_val}_{total_accounts}.csv"
        st.session_state.processed_data = df_out

        # --- RELEASE FILE ---
        ref_col_release = None
        for col in df_src.columns:
            col_clean = col.strip().upper().replace(" ", "").replace("_", "")
            if col_clean in ['REFCODE', 'REFERENCECODE', 'REFCODE']:
                ref_col_release = col
                break

        released_to_col = None
        if version == "OTS":
            for col in df_src.columns:
                col_clean = col.strip().upper().replace(" ", "").replace("_", "")
                if col_clean == 'TAGFM':
                    released_to_col = col
                    break
        else:   # For MC2, try to find a column that likely holds the released‑to information
            for col in df_src.columns:
                col_clean = col.strip().upper().replace(" ", "").replace("_", "")
                if col_clean in ['TAGFM', 'RELEASEDTO', 'TAG']:
                    released_to_col = col
                    break

        df_rel = pd.DataFrame()
        df_rel['REF CODE'] = df_src[ref_col_release] if ref_col_release else ""
        df_rel['RELEASED TO'] = df_src[released_to_col] if released_to_col else ""
        st.session_state.release_data = df_rel

        progress_bar.progress(100, text="Done!")
        elapsed = time.time() - start_time
        progress_bar.empty()
        st.success(f"✅ Processed {total_accounts} accounts in {elapsed:.2f} seconds.")
        play_completion_sound()

    # Display results
    if st.session_state.processed_data is not None:
        df_out = st.session_state.processed_data
        df_release = st.session_state.release_data
        main_fn = st.session_state.generated_filename
        rel_fn = st.session_state.release_filename
        total_accounts = len(df_out)

        _, center_box, _ = st.columns([1, 1, 1])
        with center_box:
            st.success(f"✅ Processed {total_accounts} accounts!")

        st.write("### 📥 Download Processed Files")
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            csv_main = df_out.to_csv(index=False).encode('utf-8')
            st.download_button(
                label=f"📥 Main CSV ({main_fn})",
                data=csv_main,
                file_name=main_fn,
                mime="text/csv",
                use_container_width=True
            )
        with col_dl2:
            csv_rel = df_release.to_csv(index=False).encode('utf-8')
            st.download_button(
                label=f"📊 Released to ({rel_fn})",
                data=csv_rel,
                file_name=rel_fn,
                mime="text/csv",
                use_container_width=True
            )

        st.divider()
        st.subheader("📋 Main Data Preview")
        st.dataframe(df_out.head(10))
        st.subheader("📑 Release File Preview")
        st.dataframe(df_release.head(10))


# ==========================================
# TOOL 2: FIELD RESULT (REPLICATED VBA MACRO)
# ==========================================
elif selected_tool == "Field Result":
    st.title("📊 Field Result Column Extractor")
    st.write("Upload an Excel workbook containing a sheet named **'RESULT'** to run the macro extraction.")

    excel_file = st.file_uploader("Upload Workbook (.xlsx)", type=['xlsx'], key="field_result_uploader")

    if excel_file and st.button("🔄 Process Field Result", use_container_width=True):
        start_time = time.time()
        try:
            progress_bar = st.progress(0, text="Reading Excel file...")
            xl_file = pd.ExcelFile(excel_file)
            sheet_names = xl_file.sheet_names
            target_sheet = next((sheet for sheet in sheet_names if sheet.upper() == "RESULT"), None)

            if not target_sheet:
                progress_bar.empty()
                st.error("❌ Error: A sheet named 'RESULT' was not found in this workbook.")
                st.session_state.field_result_processed = False
            else:
                progress_bar.progress(25, text="Extracting relevant columns...")
                time.sleep(0.2)

                df_source = pd.read_excel(excel_file, sheet_name=target_sheet, header=None, dtype=str)
                max_required_col_idx = 41
                if df_source.shape[1] <= max_required_col_idx:
                    for i in range(df_source.shape[1], max_required_col_idx + 1):
                        df_source[i] = ""

                df_target = pd.DataFrame()
                df_target[0] = df_source[41]
                df_target[1] = df_source[2]
                df_target[2] = df_source[3]
                df_target[3] = df_source[4]
                df_target[4] = df_source[27]
                df_target[5] = df_source[32]
                df_target[6] = df_source[36]

                progress_bar.progress(50, text="Converting dates to native objects...")
                time.sleep(0.2)

                def format_field_result_date(val):
                    if pd.isna(val):
                        return val
                    val_str = str(val).strip()
                    if re.match(r'^\d{4}-\d{2}-\d{2}( \d{2}:\d{2}:\d{2}(\.\d+)?)?$', val_str) or \
                       re.match(r'^\d{1,2}/\d{1,2}/\d{4}( \d{2}:\d{2}:\d{2}(\.\d+)?)?$', val_str):
                        try:
                            return pd.to_datetime(val_str).date()
                        except:
                            return val
                    return val

                for col in df_target.columns:
                    df_target[col] = df_target[col].apply(format_field_result_date)

                progress_bar.progress(75, text="Generating stylized Excel with mm/dd/yyyy format...")
                time.sleep(0.2)

                wb_out = openpyxl.Workbook()
                ws_out = wb_out.active
                ws_out.title = "RESULT"
                calibri_font = Font(name="Calibri", size=9)

                for r_idx, row_values in enumerate(df_target.values):
                    for c_idx, cell_value in enumerate(row_values):
                        if isinstance(cell_value, (date, datetime)):
                            formatted_text_value = cell_value.strftime('%m/%d/%Y')
                        else:
                            formatted_text_value = cell_value
                        cell = ws_out.cell(row=r_idx + 1, column=c_idx + 1, value=formatted_text_value)
                        cell.font = calibri_font

                for col in ws_out.columns:
                    max_len = 0
                    col_letter = get_column_letter(col[0].column)
                    for cell in col:
                        if cell.value is not None:
                            max_len = max(max_len, len(str(cell.value)))
                    ws_out.column_dimensions[col_letter].width = max(max_len + 3, 10)

                current_date_str = datetime.now().strftime('%m-%d-%Y')
                generated_fn = f"FIELD RESULT {current_date_str}.xlsx"

                output_buffer = io.BytesIO()
                wb_out.save(output_buffer)
                output_buffer.seek(0)

                st.session_state.field_result_buffer = output_buffer
                st.session_state.field_result_filename = generated_fn
                st.session_state.field_result_processed = True
                elapsed = time.time() - start_time
                st.session_state.field_result_elapsed = elapsed

                progress_bar.progress(100, text="Process Complete!")
                time.sleep(0.3)
                progress_bar.empty()

                st.success(f"✅ Extraction and formatting (mm/dd/yyyy) successful! (Elapsed: {elapsed:.2f}s)")
                play_completion_sound()

                preview_df = df_target.copy()
                preview_df.columns = preview_df.iloc[0]
                preview_df = preview_df.drop(preview_df.index[0]).reset_index(drop=True)
                st.divider()
                st.subheader("📋 Extracted Data Preview")
                st.dataframe(preview_df.head(15))

        except Exception as e:
            st.error(f"❌ Error processing file: {str(e)}")
            st.session_state.field_result_processed = False

    if st.session_state.field_result_processed and st.session_state.field_result_buffer is not None:
        st.download_button(
            label=f"📥 Download Extracted File ({st.session_state.field_result_filename})",
            data=st.session_state.field_result_buffer,
            file_name=st.session_state.field_result_filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="field_result_download"
        )


# ==========================================
# TOOL 3: E-SIGN FIXER
# ==========================================
elif selected_tool == "E-SIGN FIXER":
    st.title("✒️ E-SIGN FIXER")
    st.write("Upload an image page containing multiple sheet signatures to crop them out into isolated transparent files.")

    TARGET_SIZE = 250
    uploaded_file = st.file_uploader("Upload Image", type=["png", "jpg", "jpeg"])

    if uploaded_file:
        try:
            start_time = time.time()
            progress_bar = st.progress(0, text="Loading image...")
            image = Image.open(uploaded_file).convert("RGB")
            img = np.array(image)

            with st.expander("🖼️ Show Original Image", expanded=False):
                st.image(img, width=350)

            progress_bar.progress(25, text="Processing...")
            time.sleep(0.25)

            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            _, thresh = cv2.threshold(gray, 220, 255, cv2.THRESH_BINARY_INV)
            kernel = np.ones((3, 3), np.uint8)
            thresh = cv2.dilate(thresh, kernel, iterations=1)

            progress_bar.progress(50, text="Extracting signatures...")
            time.sleep(0.25)

            contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            contours = sorted(contours, key=lambda c: (cv2.boundingRect(c)[1], cv2.boundingRect(c)[0]))

            extracted_signatures = []

            progress_bar.progress(75, text="Cropping...")
            time.sleep(0.25)

            for cnt in contours:
                area = cv2.contourArea(cnt)
                if area < 1000:
                    continue
                x, y, w, h = cv2.boundingRect(cnt)
                padding = 15
                x1 = max(0, x - padding)
                y1 = max(0, y - padding)
                x2 = min(img.shape[1], x + w + padding)
                y2 = min(img.shape[0], y + h + padding)
                crop = img[y1:y2, x1:x2]
                mask = np.any(crop < 240, axis=2).astype(np.uint8) * 255
                ys, xs = np.where(mask > 0)
                if len(xs) == 0 or len(ys) == 0:
                    continue
                crop = crop[ys.min():ys.max() + 1, xs.min():xs.max() + 1]
                mask = mask[ys.min():ys.max() + 1, xs.min():xs.max() + 1]
                gray_crop = cv2.cvtColor(crop, cv2.COLOR_RGB2GRAY)
                rgba = np.zeros((gray_crop.shape[0], gray_crop.shape[1], 4), dtype=np.uint8)
                rgba[:, :, 0:3] = gray_crop[:, :, np.newaxis]
                rgba[:, :, 3] = mask
                h_sig, w_sig = rgba.shape[:2]
                max_dim = max(h_sig, w_sig)
                square_canvas = np.zeros((max_dim, max_dim, 4), dtype=np.uint8)
                y_offset = (max_dim - h_sig) // 2
                x_offset = (max_dim - w_sig) // 2
                square_canvas[y_offset:y_offset+h_sig, x_offset:x_offset+w_sig] = rgba
                final_sig = cv2.resize(square_canvas, (TARGET_SIZE, TARGET_SIZE), interpolation=cv2.INTER_AREA)
                extracted_signatures.append(final_sig)

            progress_bar.progress(90, text="Generating files...")

            if extracted_signatures:
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for idx, sig in enumerate(extracted_signatures, start=1):
                        filename = f"signature_{idx}.PNG"
                        img_byte_arr = io.BytesIO()
                        Image.fromarray(sig).save(img_byte_arr, format='PNG')
                        zipf.writestr(filename, img_byte_arr.getvalue())

                doc = Document()
                doc.add_heading('Extracted E-Signatures', 0)
                for idx, sig in enumerate(extracted_signatures, start=1):
                    img_byte_arr = io.BytesIO()
                    Image.fromarray(sig).save(img_byte_arr, format='PNG')
                    img_byte_arr.seek(0)
                    doc.add_paragraph(f'Signature {idx}:')
                    doc.add_picture(img_byte_arr, width=Inches(2.5))

                word_buffer = BytesIO()
                doc.save(word_buffer)

                progress_bar.progress(100, text="Complete!")
                elapsed = time.time() - start_time
                st.success(f"✅ Extracted {len(extracted_signatures)} signatures in {elapsed:.2f} seconds.")
                play_completion_sound()

                c1, c2 = st.columns(2)
                with c1:
                    st.download_button("📦 ZIP Images", zip_buffer.getvalue(), "signatures.zip", "application/zip", use_container_width=True)
                with c2:
                    st.download_button("📝 Word Document", word_buffer.getvalue(), "signatures.docx", use_container_width=True)

                st.divider()
                cols = st.columns(4)
                for idx, sig in enumerate(extracted_signatures):
                    cols[idx % 4].image(sig, caption=f"Sig {idx+1}")
            else:
                progress_bar.empty()
                st.warning("No signatures were detected. Try adjusting the image or threshold.")

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
